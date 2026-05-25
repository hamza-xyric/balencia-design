#!/usr/bin/env python3
"""
Update Xyric_Deployment_Request_Form.docx with completed content.
Run: python scripts/update-deployment-form-docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_PATH = Path(__file__).parent.parent / "Xyric_Deployment_Request_Form.docx"
OUTPUT_PATH = Path(__file__).parent.parent / "Xyric_Deployment_Request_Form_UPDATED.docx"

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.color.rgb = None  # Default
    p.space_after = Pt(6)
    return p

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p

def main():
    doc = Document(DOCX_PATH)
    
    # Find and update key sections by iterating (simplified - we add system design at end)
    # For full replacement, we'd need to modify existing paragraphs. Instead, append.
    
    # Add System Design section at end (before sectPr)
    add_heading(doc, "11. Complete System Design", level=1)
    add_paragraph(doc, "See Xyric_Deployment_Request_Form_COMPLETED.md for full system design.")
    add_paragraph(doc, "")
    add_paragraph(doc, "Architecture: Client (Next.js) → API (Express) → Services (LangChain, WHOOP) → PostgreSQL.")
    add_paragraph(doc, "Components: Auth (JWT), RBAC, WHOOP integration, Voice calls, Chat, Wellbeing, Activity Status.")
    
    doc.save(OUTPUT_PATH)
    print(f"Updated docx saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
