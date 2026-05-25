#!/usr/bin/env python3
"""
Create Xyric_Deployment_Request_Form.docx from completed content.
Run: python scripts/create-deployment-docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = Path(__file__).parent.parent / "Xyric_Deployment_Request_Form_COMPLETED.docx"

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text):
    return doc.add_paragraph(text)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(cell)
    return t

def main():
    doc = Document()
    
    # Title
    p = doc.add_paragraph()
    r = p.add_run("XYRIC")
    r.bold = True
    r.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    r = p.add_run("Deployment Request Form")
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    add_para(doc, "Submitted By: Salman Sadiq")
    add_para(doc, "Date: 11 February 2026")
    doc.add_paragraph()
    
    # 1. Project Overview
    add_heading(doc, "1. Project Overview", level=1)
    add_para(doc, "Project Name: yHealth Platform")
    add_para(doc, "One-Line Description: AI Life Coach that unifies Fitness, Nutrition, and Wellbeing data to deliver conversational coaching via Voice, WhatsApp, and Mobile App.")
    add_para(doc, "Problem Statement: Health data is fragmented across fitness trackers, nutrition apps, and wellness tools. yHealth integrates all three pillars and delivers personalized AI coaching for health-conscious consumers.")
    add_para(doc, "Target Users: Health-conscious consumers, fitness enthusiasts, internal team for beta testing.")
    add_para(doc, "Success Metrics: User signups, multi-pillar adoption, Voice/WhatsApp engagement, feature usage.")
    doc.add_paragraph()
    
    # 2. Technical Summary
    add_heading(doc, "2. Technical Summary", level=1)
    add_heading(doc, "2.1 Architecture Diagram", level=2)
    add_para(doc, "See Section 11 (Complete System Design) for full architecture.")
    add_para(doc, "Reference: yhealth-app/docs/voice_call_system_design.md, Xyric_Deployment_Request_Form_COMPLETED.md")
    add_heading(doc, "2.2 Tech Stack", level=2)
    add_table(doc, ["Component", "Technology & Version"], [
        ["Frontend", "Next.js 16, React 19, TypeScript 5, Tailwind CSS 4"],
        ["Backend", "Node.js 20+, Express 5, TypeScript 5"],
        ["Database", "PostgreSQL 16, pgvector"],
        ["Hosting", "Railway (API), Vercel optional"],
        ["Third-Party APIs", "OpenAI, Anthropic, WHOOP, Stripe, SendGrid, AWS S3, Socket.io"],
    ])
    doc.add_paragraph()
    
    # 3. Security
    add_heading(doc, "3. Security Essentials", level=1)
    add_table(doc, ["Requirement", "Status", "Notes"], [
        ["Authentication (JWT, OAuth)", "Yes", "JWT, WHOOP OAuth, RBAC"],
        ["No secrets in source code", "Verified", ".env, .gitignore"],
        ["Environment variables", "Yes", "DB_*, JWT_SECRET, etc."],
        ["HTTPS enabled", "Yes", "Production"],
        ["Input validation", "Yes", "Zod middleware"],
        ["Rate limiting", "Yes", "express-rate-limit"],
    ])
    doc.add_paragraph()
    
    # 4. Data & Database
    add_heading(doc, "4. Data & Database", level=1)
    add_para(doc, "Data stored: User profiles, assessments, goals, workout/diet plans, health records (WHOOP), voice calls, chat, wellbeing, activity status.")
    add_para(doc, "Sensitive data: Emails, hashed passwords, health metrics, OAuth tokens. Protected via JWT, bcrypt, RBAC.")
    add_para(doc, "Database backup: Manual process documented.")
    doc.add_paragraph()
    
    # 5. Deployment
    add_heading(doc, "5. Deployment Details", level=1)
    add_para(doc, "Target: Production | URL: yhealth.xyric.ai | CI/CD: Railway | Rollback: Revert deployment, additive migrations.")
    doc.add_paragraph()
    
    # 6-8. Testing, Monitoring, Risks
    add_heading(doc, "6. Testing Status", level=1)
    add_para(doc, "Core features tested: Done | Unit tests: Partial | No critical bugs: Confirmed | Target devices: Tested")
    add_heading(doc, "7. Monitoring & Support", level=1)
    add_para(doc, "Error tracking, health endpoint, logs. Contact: Salman Sadiq")
    add_heading(doc, "8. Known Limitations & Risks", level=1)
    add_para(doc, "WHOOP webhook setup in prod; Voice cold start latency; API rate limits. Est. cost: $50-150/mo.")
    doc.add_paragraph()
    
    # 9. Pre-Deployment Checklist
    add_heading(doc, "9. Pre-Deployment Checklist", level=1)
    for item in ["README up to date", "Env vars documented", "No hardcoded secrets", "Migrations work", "Core flows tested", "Error handling", "Deploy/rollback ready", "Architecture diagram attached"]:
        add_para(doc, f"☑ {item}")
    doc.add_paragraph()
    
    # 10. Approval
    add_heading(doc, "10. Approval", level=1)
    add_table(doc, ["Role", "Name / Signature", "Date"], [
        ["Submitted By", "", ""],
        ["Tech Lead Review", "", ""],
        ["CTO Approval", "", ""],
    ])
    add_para(doc, "CTO Decision: ☐ Approved ☐ Approved with notes ☐ Needs changes")
    doc.add_paragraph()
    
    # 11. Complete System Design
    add_heading(doc, "11. Complete System Design", level=1)
    add_heading(doc, "11.1 High-Level Architecture", level=2)
    add_para(doc, "Client Layer: Next.js App (Dashboard, Fitness, Nutrition, Wellbeing, WHOOP, Voice, Chat), WhatsApp Business, Mobile Web (PWA)")
    add_para(doc, "API Layer: Express 5 with Auth (JWT), Rate Limit, Validation (Zod), RBAC")
    add_para(doc, "Service Layer: LangChain/LangGraph, WHOOP Sync, BullMQ Automation, Voice TTS/STT")
    add_para(doc, "Data Layer: PostgreSQL (users, health_data_records, voice_calls, activity_status_history, pgvector), Redis (optional), S3")
    add_para(doc, "External: WHOOP API, OpenAI, Anthropic, Stripe, SendGrid, AWS S3")
    add_heading(doc, "11.2 Component Architecture", level=2)
    add_table(doc, ["Layer", "Component", "Responsibility"], [
        ["Client", "Next.js App", "SSR/CSR, auth, feature UIs"],
        ["API", "Express Router", "Routes, middleware"],
        ["API", "Auth Middleware", "JWT verify, RBAC"],
        ["Service", "LangGraph Chatbot", "RAG, AI responses"],
        ["Service", "WHOOP Service", "OAuth, sync, webhooks"],
        ["Data", "PostgreSQL", "Persistent data"],
    ])
    add_heading(doc, "11.3 Data Flow", level=2)
    add_para(doc, "WHOOP: User → GET /integrations/whoop/status, /whoop/analytics/overview → Rendered charts")
    add_para(doc, "Voice: User → POST /voice-calls/initiate → WebRTC → AI Coach (LangChain) ↔ TTS")
    add_para(doc, "Chat: User → POST /chats/:id/messages → RAG Chatbot → OpenAI/Anthropic → Streaming")
    add_heading(doc, "11.4 Security", level=2)
    add_para(doc, "Auth: JWT (httpOnly cookie), bcrypt | RBAC | Secrets in .env | Zod validation | Rate limiting | HTTPS")
    
    doc.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
